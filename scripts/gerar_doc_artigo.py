# -*- coding: utf-8 -*-
"""Converte output/<pasta>/artigo.html em um .docx editavel com estilos nativos do Word."""
import os, io, re, sys
from html.parser import HTMLParser

import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

BLOCOS = {"h1", "h2", "h3", "h4", "h5", "p", "li", "hr", "img", "blockquote"}


class Bloco:
    def __init__(self, kind):
        self.kind = kind
        self.runs = []          # (texto, bold, italic, underline, href)
        self.img = None         # (src, alt)


class Parser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocos = []
        self.atual = None
        self.fmt = {"b": 0, "i": 0, "u": 0}
        self.href = None
        self.em_quote = False

    def _abre(self, kind):
        self.atual = Bloco(kind)
        self.blocos.append(self.atual)

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        if tag in ("h1", "h2", "h3", "h4", "h5"):
            self._abre(tag)
        elif tag == "p":
            self._abre("quote" if self.em_quote else "p")
        elif tag == "li":
            self._abre("li")
        elif tag == "blockquote":
            self.em_quote = True
        elif tag == "hr":
            self._abre("hr"); self.atual = None
        elif tag == "img":
            b = Bloco("img"); b.img = (a.get("src", ""), a.get("alt", ""))
            self.blocos.append(b); self.atual = None
        elif tag in ("strong", "b"):
            self.fmt["b"] += 1
        elif tag in ("em", "i"):
            self.fmt["i"] += 1
        elif tag == "u":
            self.fmt["u"] += 1
        elif tag == "a":
            self.href = a.get("href")

    def handle_endtag(self, tag):
        if tag == "blockquote":
            self.em_quote = False
        elif tag in ("strong", "b"):
            self.fmt["b"] = max(0, self.fmt["b"] - 1)
        elif tag in ("em", "i"):
            self.fmt["i"] = max(0, self.fmt["i"] - 1)
        elif tag == "u":
            self.fmt["u"] = max(0, self.fmt["u"] - 1)
        elif tag == "a":
            self.href = None
        elif tag in ("h1", "h2", "h3", "h4", "h5", "p", "li"):
            self.atual = None

    def handle_data(self, data):
        if self.atual is None:
            return
        txt = re.sub(r"\s+", " ", data)
        if not txt.strip() and not self.atual.runs:
            return
        self.atual.runs.append((txt, self.fmt["b"] > 0, self.fmt["i"] > 0,
                                self.fmt["u"] > 0, self.href))


def add_hyperlink(par, texto, url, bold=False):
    """Insere um hyperlink real (python-docx nao tem API nativa)."""
    part = par.part
    r_id = part.relate_to(url,
                          docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                          is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = texto
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    link.append(r)
    par._p.append(link)


def linha_horizontal(par):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "BFBFBF")):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


def sombrear(celula, hexcor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcor)
    celula._tc.get_or_add_tcPr().append(shd)


def montar(cfg):
    html = io.open(cfg["arquivo"], encoding="utf-8").read()
    p = Parser(); p.feed(html)

    doc = Document()

    # pagina e fonte base
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.6)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    titulo = "".join(t for t, *_ in p.blocos[0].runs) if p.blocos else cfg["slug"]

    # cabecalho de identificacao
    k = doc.add_paragraph()
    rk = k.add_run(cfg["veiculo"].upper() + "  |  " + cfg["tom"].upper())
    rk.bold = True; rk.font.size = Pt(8.5)
    rk.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    k.paragraph_format.space_after = Pt(2)

    doc.add_heading(titulo, level=0)

    # ficha
    campos = [
        ("Situação", cfg["status"]),
        ("Título SEO", cfg["seo"]),
        ("Meta description", cfg["meta"]),
        ("Slug", cfg["slug"]),
        ("Origem", "WakeCast episódio 5, com Diego Santos (Wake) e Ian (metaKosmos)"),
    ]
    tab = doc.add_table(rows=0, cols=2)
    tab.style = "Table Grid"
    tab.alignment = WD_TABLE_ALIGNMENT.LEFT
    for rot, val in campos:
        row = tab.add_row()
        c0, c1 = row.cells
        sombrear(c0, "F2F2F2")
        r0 = c0.paragraphs[0].add_run(rot)
        r0.bold = True; r0.font.size = Pt(9)
        r1 = c1.paragraphs[0].add_run(val)
        r1.font.size = Pt(9)
        c0.width = Cm(3.6); c1.width = Cm(12.4)
    for row in tab.rows:
        for c in row.cells:
            for par in c.paragraphs:
                par.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    aviso = doc.add_paragraph()
    ra = aviso.add_run(cfg["aviso"])
    ra.italic = True; ra.font.size = Pt(9)
    ra.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    linha_horizontal(aviso)
    doc.add_paragraph()

    # corpo
    mapa_h = {"h2": 1, "h3": 2, "h4": 3, "h5": 4}
    for b in p.blocos[1:]:
        if b.kind == "hr":
            par = doc.add_paragraph(); linha_horizontal(par); continue

        if b.kind == "img":
            src, alt = b.img
            par = doc.add_paragraph()
            r = par.add_run("[IMAGEM]  " + alt)
            r.bold = True; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
            r2 = par.add_run("\n" + src)
            r2.font.size = Pt(8); r2.font.name = "Consolas"
            r2.font.color.rgb = RGBColor(0xA6, 0xA6, 0xA6)
            linha_horizontal(par)
            continue

        if b.kind in mapa_h:
            par = doc.add_heading("", level=mapa_h[b.kind])
        elif b.kind == "li":
            par = doc.add_paragraph(style="List Bullet")
        elif b.kind == "quote":
            par = doc.add_paragraph(style="Intense Quote")
        else:
            par = doc.add_paragraph()

        for texto, bold, ital, undl, href in b.runs:
            if href:
                add_hyperlink(par, texto, href, bold=bold)
                continue
            # destaca placeholders do tipo [INSERIR_XXX]
            for pedaco in re.split(r"(\[INSERIR_[A-Z_]+\])", texto):
                if not pedaco:
                    continue
                r = par.add_run(pedaco)
                r.bold = bold; r.italic = ital; r.underline = undl
                if pedaco.startswith("[INSERIR_"):
                    r.bold = True
                    r.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

    doc.core_properties.title = titulo
    doc.core_properties.author = "metaKosmos"
    doc.core_properties.comments = cfg["status"]

    doc.save(cfg["saida"])
    return titulo, len(p.blocos)


ARTIGOS = [
    {
        "arquivo": os.path.join(BASE, "output", "provador-virtual-ecommerce-moda", "artigo.html"),
        "saida": os.path.join(BASE, "output", "1 - Artigo metaKosmos - Provador virtual no e-commerce de moda.docx"),
        "veiculo": "Blog metaKosmos",
        "tom": "Comercial",
        "status": "Rascunho já criado no Payload (post ID 68). Aguardando aprovação.",
        "seo": "Provador virtual no e-commerce de moda: venda mais (50 caracteres)",
        "meta": "Provador virtual de moda com IA: recupere parte dos 98 de cada 100 que abandonam, "
                "reduza devoluções em até 61% e corte até 90% do shooting. (139 caracteres)",
        "slug": "provador-virtual-ecommerce-moda",
        "aviso": "As imagens estão marcadas como [IMAGEM] com a descrição e o nome do arquivo, "
                 "porque os arquivos vivem na biblioteca de mídia do site. Os links estão ativos e "
                 "já carregam os parâmetros de rastreamento. Dois números dependem de conferência: "
                 "o percentual exato da Gregory e o tamanho do estudo State of Immersive and "
                 "Agentic Commerce. Nenhum valor monetário foi publicado.",
    },
    {
        "arquivo": os.path.join(BASE, "output", "_WAKE-nao-publicar-no-payload",
                                "inovacao-ecommerce-processo", "artigo.html"),
        "saida": os.path.join(BASE, "output", "2 - Artigo Wake - Inovacao no e-commerce virou processo.docx"),
        "veiculo": "Blog da Wake",
        "tom": "Neutro / editorial de plataforma",
        "status": "Rascunho para o time da Wake revisar e publicar. Não vai para o site da metaKosmos.",
        "seo": "Inovação no e-commerce: virou processo, não projeto (51 caracteres)",
        "meta": "3D, realidade aumentada e provador virtual ficaram acessíveis. Como estruturar um "
                "teste com métrica clara e sem travar a operação. (130 caracteres)",
        "slug": "inovacao-ecommerce-processo",
        "aviso": "Texto escrito na perspectiva da Wake, sem citar a metaKosmos nominalmente, para "
                 "preservar a autoridade neutra de um editorial de plataforma. Há um marcador "
                 "destacado em amarelo, [INSERIR_CTA_WAKE], que o time da Wake precisa preencher. "
                 "Nenhuma URL da Wake foi inventada e o artigo foi entregue sem imagens, porque não "
                 "tenho acesso à biblioteca de mídia deles.",
    },
]

if __name__ == "__main__":
    for cfg in ARTIGOS:
        t, n = montar(cfg)
        print("[OK] %s\n     %d blocos | %s" % (os.path.basename(cfg["saida"]), n, t))
